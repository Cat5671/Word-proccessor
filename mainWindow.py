import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QTextEdit, QSizePolicy, QMenu, QAction, \
    QInputDialog, QMessageBox, QLabel
from bs4 import BeautifulSoup
from PyQt5.QtCore import Qt, QUrl, QRegExp
from PyQt5 import QtCore
from PyQt5.QtGui import QColor, QFont, QDesktopServices, QTextCursor, QImage, QTextImageFormat, QTextCharFormat, QBrush, \
    QTextBlockFormat
from docx.oxml import OxmlElement, ns

from layoutWordProccessor1 import Ui_WordProcessor
from docx import Document
from docx2pdf import convert
import os
import tempfile
from docx.shared import Pt, RGBColor
from PyPDF2 import PdfReader
from docx.oxml.ns import qn
import base64


class Sheet(QTextEdit):

    def __init__(self, sheets_layout, previous_sheet=None):
        super().__init__()
        self.previous_sheet = previous_sheet
        self.sheets_layout = sheets_layout
        size_policy = QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Expanding)
        size_policy.setHorizontalStretch(0)
        size_policy.setVerticalStretch(0)
        size_policy.setHeightForWidth(self.sizePolicy().hasHeightForWidth())
        self.setSizePolicy(size_policy)
        self.setMinimumSize(QtCore.QSize(990, 1400))
        self.setMaximumSize(QtCore.QSize(990, 1400))
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setStyleSheet("background-color: rgb(255, 255, 255);\n"
                           "QMenuBar::item:selected {\n"
                           "background-color: #AAAAAA;\n"
                           " }")

        sheets_layout.addWidget(self)
        self.current_font = "MS Shell Dlg 2"
        self.current_color = "Black"
        self.current_size = "8"
        self.setStyleSheet("""
                    QTextEdit {
                        background-color: #FFFFFF;
                        padding-left: 144px;     /* Отступ слева (1.27 см) */
                        padding-right: 72px;    /* Отступ справа (1.27 см) */
                        padding-top: 96px;      /* Отступ сверху (1.27 см) */
                        padding-bottom: 96px;   /* Отступ снизу (1.27 см) */              
                    }
                """)

        self.page_label = QLabel(f"{self.sheets_layout.indexOf(self) + 1}", self)
        self.page_label.setStyleSheet("color: black; font-size: 26px; background-color: white;")
        self.page_label.setAttribute(Qt.WA_TransparentForMouseEvents)
        self.page_label_position = "center"
        self.page_label.resize(100, 30)
        self.page_label.move(463, 1260)
        self.page_label.hide()
        self.is_page_numbering_enabled = False

        self.f = -1
        self.sheet1 = None
        self.textChanged.connect(self.check_and_delete_if_empty)
        self.textChanged.connect(lambda: self.check_text_height(sheets_layout))
        self.cursorPositionChanged.connect(self.apply_styles_to_current_sheet)

    def mousePressEvent(self, event):
        if event.button() == 1:
            if event.modifiers() == Qt.ControlModifier:
                cursor = self.cursorForPosition(event.pos())
                if cursor.charFormat().anchorHref():
                    url = cursor.charFormat().anchorHref()
                    QDesktopServices.openUrl(QUrl(url))
        super().mousePressEvent(event)

    def toggle_page_numbers(self, enable):
        self.is_page_numbering_enabled = enable
        if enable:
            self.page_label.show()
        else:
            self.page_label.hide()

    def check_text_height(self, sheets_layout):
        check = self.cursorRect(self.textCursor()).y() + self.cursorRect().height() * 2
        check2 = self.document().size().height()

        if check2 > 1208 and self.sheet1 is None:
            self.sheet1 = Sheet(sheets_layout, previous_sheet=self)
            self.text = self.toPlainText().split("\n")
            self.sheet1.insertPlainText(self.text.pop())

        elif check2 > 1208:
            if check > 1208:
                self.f -= 1
            self.text = self.toPlainText().split("\n")

            self.sheet1.insertPlainText("\n" + self.text[len(self.text) + self.f])

        if check > 1208 and self.sheet1 is None:
            self.sheet1 = Sheet(sheets_layout, previous_sheet=self)
            self.apply_styles_to_sheet(self.sheet1)
            self.move_cursor(self.sheet1)

        elif check > 1208:
            self.apply_styles_to_sheet(self.sheet1)
            self.move_cursor(self.sheet1)

    def check_and_delete_if_empty(self):
        text = self.toPlainText().strip()
        cursor = self.textCursor()
        if not text and not cursor.hasSelection():
            self.remove_sheet()

    def remove_sheet(self):
        if self.previous_sheet is None:
            print(2)
            return
        self.sheets_layout.removeWidget(self)

        if self.previous_sheet is not None:
            self.previous_sheet.sheet1 = None

        self.deleteLater()

    def apply_styles_to_current_sheet(self):
        self.set_font()
        self.set_size()
        self.set_color()
        self.set_underline(self.fontUnderline())
        self.set_bold(self.fontWeight())
        self.set_italic(self.fontItalic())

    def apply_styles_to_sheet(self, sheet):
        sheet.set_font(self.current_font)
        sheet.set_size(self.current_size)
        sheet.set_color(self.current_color)
        sheet.set_underline(self.fontUnderline())
        sheet.set_bold(self.fontWeight())
        sheet.set_italic(self.fontItalic())

    def move_cursor(self, sheet):
        self.cursor = sheet.textCursor()
        self.cursor.setPosition(0)
        sheet.setFocus()
        sheet.setTextCursor(self.cursor)

    def set_font(self, font="mS Shell Dlg 2"):

        if self.textCursor().selectionEnd() == self.textCursor().selectionStart():
            if font != "mS Shell Dlg 2":
                self.current_font = font

            self.setFontFamily(self.current_font)
            if self.sheet1 is not None:
                self.sheet1.set_font(self.current_font)

        elif font != "mS Shell Dlg 2":  # условие для форматирования выделенного текста
            self.current_font = font
            self.setFontFamily(self.current_font)
            if self.sheet1 is not None:
                self.sheet1.set_font(self.current_font)

    def set_size(self, font_size="7"):
        if self.textCursor().selectionEnd() == self.textCursor().selectionStart():
            if font_size != "7":
                self.current_size = font_size

            self.setFontPointSize(int(self.current_size))
            if self.sheet1 is not None:
                self.sheet1.set_size(self.current_size)

        elif font_size != "7":
            self.current_size = font_size
            self.setFontPointSize(int(self.current_size))
            if self.sheet1 is not None:
                self.sheet1.set_size(self.current_size)

    def set_color(self, font_color="black"):
        if self.textCursor().selectionEnd() == self.textCursor().selectionStart():
            if font_color != "black":
                self.current_color = font_color

                self.setTextColor(QColor(self.current_color))
                if self.sheet1 is not None:
                    self.sheet1.set_color(self.current_color)

        elif font_color != "black":
            self.current_color = font_color
            self.setTextColor(QColor(self.current_color))
            if self.sheet1 is not None:
                self.sheet1.set_color(self.current_color)

    def count_u(self):
        if self.fontUnderline():
            self.set_underline(False)
        else:
            self.set_underline(True)

    def count_b(self):
        if self.fontWeight():
            self.set_bold(False)
        else:
            self.set_bold(QFont.Bold)

    def count_i(self):
        if self.fontItalic():
            self.set_italic(False)
        else:
            self.set_italic(True)

    def set_underline(self, c):
        self.setFontUnderline(c)
        if self.sheet1 is not None:
            self.sheet1.set_underline(c)

    def set_bold(self, c):
        self.setFontWeight(c)
        if self.sheet1 is not None:
            self.sheet1.set_bold(c)

    def set_italic(self, c):
        self.setFontItalic(c)
        if self.sheet1 is not None:
            self.sheet1.set_italic(c)

    def add_hyperlink(self):
        cursor = self.textCursor()
        selected_text = cursor.selectedText()

        if not selected_text:
            QMessageBox.warning(self, "Ошибка", "Выберите текст для добавления гиперссылки.")
            return

        url, ok = QInputDialog.getText(self, "Добавить гиперссылку", "Введите URL:")
        if ok and url:
            char_format = cursor.charFormat()
            html_link = f'<a href="{url}" style="color: blue; ' \
                        f'font-size: {char_format.font().pointSize()}pt; ' \
                        f'font-family: {char_format.font().family()}; ' \
                        f'font-weight: {"bold" if char_format.font().bold() else "normal"}; ' \
                        f'font-style: {"italic" if char_format.font().italic() else "normal"}; ' \
                        f'text-decoration: underline;">{selected_text}</a>'
            cursor.insertHtml(html_link)

    def remove_hyperlink(self):
        cursor = self.textCursor()
        char_format = cursor.charFormat()

        if not char_format.anchorHref():
            QMessageBox.warning(self, "Ошибка", "Гиперссылка не найдена.")
            return

        target_url = char_format.anchorHref()

        cursor.beginEditBlock()
        cursor.movePosition(QTextCursor.Start)
        while not cursor.atEnd():
            cursor.movePosition(QTextCursor.NextCharacter, QTextCursor.KeepAnchor)
            char_format = cursor.charFormat()
            if char_format.anchorHref() == target_url:
                char_format.setAnchor(False)
                char_format.setAnchorHref("")
                cursor.setCharFormat(char_format)
        cursor.endEditBlock()

    def insert_image(self):
        file_dialog = QFileDialog(self)
        file_dialog.setNameFilter("Изображения (*.png *.jpg *.jpeg *.bmp *.gif)")
        file_dialog.setFileMode(QFileDialog.ExistingFile)
        if file_dialog.exec():
            image_path = file_dialog.selectedFiles()[0]
            image = QImage(image_path)

            if image.isNull():
                QMessageBox.warning(self, "Ошибка", "Не удалось загрузить изображение.")
                return

            max_width = 500
            if image.width() > max_width:
                image = image.scaledToWidth(max_width, Qt.SmoothTransformation)

            cursor = self.textCursor()
            image_format = QTextImageFormat()
            image_format.setName(image_path)
            image_format.setWidth(image.width())
            image_format.setHeight(image.height())
            cursor.insertImage(image_format)



    def contextMenuEvent(self, event):
        context_menu = QMenu(self)
        copy_action = QAction("Копировать", self)
        insert_action = QAction("Вставить", self)
        hyperlink_action = QAction("Добавить гиперссылку", self)
        remove_hyperlink_action = QAction("Удалить гиперссылку", self)
        set_line_spacing_action = QAction("Изменить межстрочный интервал", self)
        insert_image_action = QAction("Вставить изображение", self)

        context_menu.addAction(copy_action)
        context_menu.addAction(insert_action)
        context_menu.addAction(hyperlink_action)
        context_menu.addAction(set_line_spacing_action)
        context_menu.addAction(insert_image_action)

        cursor = self.textCursor()
        char_format = cursor.charFormat()
        if char_format.anchorHref():
            context_menu.addAction(remove_hyperlink_action)

        copy_action.triggered.connect(self.copy)
        insert_action.triggered.connect(self.paste)
        hyperlink_action.triggered.connect(self.add_hyperlink)
        remove_hyperlink_action.triggered.connect(self.remove_hyperlink)
        set_line_spacing_action.triggered.connect(self.set_line_spacing)
        insert_image_action.triggered.connect(self.insert_image)

        context_menu.exec(event.globalPos())

    def set_line_spacing(self):
        cursor = self.textCursor()
        if not cursor.hasSelection():
            QMessageBox.warning(self, "Ошибка", "Выделите текст, чтобы задать межстрочный интервал.")
            return

        spacing, ok = QInputDialog.getDouble(
            self,
            "Выставить межстрочный интервал",
            "Введите размер межстрочного интервала (например, 1.0):",
            1.0, 0.5, 5.0, 1
        )
        if ok:
            block_format = cursor.blockFormat()
            block_format.setLineHeight(int(spacing * 100), QTextBlockFormat.ProportionalHeight)
            cursor.mergeBlockFormat(block_format)
            self.setTextCursor(cursor)

    def iterate_sheets(self):
        current_sheet = self
        while current_sheet:
            yield current_sheet
            current_sheet = current_sheet.sheet1

    def find_word(self, word):

        cursor = self.textCursor()
        cursor.setPosition(0)
        format = QTextCharFormat()
        format.setBackground(QBrush(QColor("white")))
        cursor.movePosition(QTextCursor.End, 1)
        cursor.mergeCharFormat(format)

        if word != "":
            cursor = self.textCursor()
            format = QTextCharFormat()
            format.setBackground(QBrush(QColor("yellow")))
            regex = QtCore.QRegExp(word)
            index = regex.indexIn(self.toPlainText(), 0)

            while index != -1:
                cursor.setPosition(index)
                for _ in word:
                    cursor.movePosition(QTextCursor.Right, 1)
                cursor.mergeCharFormat(format)
                pos = index + regex.matchedLength()
                index = regex.indexIn(self.toPlainText(), pos)

        if self.sheet1 is not None:
            self.sheet1.find_word(word)

    def replace_word(self, word, replaced_word):
        if word != "":
            cursor = self.textCursor()
            index = QRegExp(word).indexIn(self.toPlainText(), 0)

            if index != -1:
                cursor.setPosition(index)
                for _ in word:
                    cursor.movePosition(QTextCursor.Right, 1)
                cursor.removeSelectedText()
                cursor.insertText(replaced_word)
            else:
                if self.sheet1 is not None:
                    self.sheet1.replace_word(word, replaced_word)


class WordProcessor(QMainWindow, Ui_WordProcessor):
    def __init__(self):
        super().__init__()
        self.setup_ui(self)
        self.resize(800, 600)
        self.current_file_path = None

        self.sheet = Sheet(self.sheets_layout)
        self.fonts.currentTextChanged.connect(self.sheet.set_font)
        self.font_size.currentTextChanged.connect(self.sheet.set_size)
        self.font_color.currentTextChanged.connect(self.sheet.set_color)
        self.underline_text_button.clicked.connect(self.sheet.count_u)
        self.italic_text_button.clicked.connect(self.sheet.count_i)
        self.bold_text_button.clicked.connect(self.sheet.count_b)
        self.button_search_word.clicked.connect(lambda: self.sheet.find_word(self.word_search_field.text()))
        self.button_replace_field.clicked.connect(lambda: self.sheet.replace_word(self.word_search_field.text(),
                                                                            self.word_replace_field.text()))

        self.is_page_numbering_enabled = False
        self.numbering_button.setCheckable(True)
        self.numbering_button.clicked.connect(self.toggle_page_numbers)

        self.save_document_action.triggered.connect(self.save_fast)
        self.save_document_where_action.triggered.connect(self.save_as_docx)
        self.save_document_where_pdf_action.triggered.connect(self.pdf_to_path)
        self.open_document_action.triggered.connect(self.open_document)
        self.add_document_action.triggered.connect(self.add_new_document)

    def add_new_document(self):
        self.current_file_path = None
        self.sheet.setPlainText("")

    def toggle_page_numbers(self):
        self.is_page_numbering_enabled = self.numbering_button.isChecked()

        for i in range(self.sheets_layout.count()):
            sheet = self.sheets_layout.itemAt(i).widget()
            if isinstance(sheet, Sheet):
                sheet.toggle_page_numbers(self.is_page_numbering_enabled)

    def save_fast(self):
        if self.current_file_path:
            self.save_to_path(self.current_file_path)
        else:
            self.save_as_docx()

    def save_as_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить файл как...", "", "Документ Word (*.docx)")
        if file_path:
            self.current_file_path = file_path
            self.save_to_path(file_path)

    def save_to_path(self, file_path):
        try:
            doc = Document()

            doc.styles['Normal'].font.name = 'Calibri'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

            html_content = ""
            for sheet in self.sheet.iterate_sheets():
                html_content += sheet.toHtml()

            print(html_content)
            soup = BeautifulSoup(html_content, "html.parser")

            for line in soup.find_all("p"):
                paragraph = doc.add_paragraph()
                for element in line.contents:
                    if element.name == "span":
                        run = paragraph.add_run(element.text)
                        style = element.attrs.get("style", "")
                        if "font-size" in style:
                            size = int(style.split("font-size:")[1].split("pt")[0])
                            run.font.size = Pt(size)
                        if "font-family" in style:
                            font_name = style.split("font-family:")[1].split(";")[0].replace("'", "")
                            run.font.name = font_name
                        if "font-weight" in style and "bold" in style.lower():
                            run.bold = True
                        if "font-style" in style and "italic" in style.lower():
                            run.italic = True
                        if "text-decoration" in style and "underline" in style.lower():
                            run.underline = True
                        if "color" in style:
                            color_value = style.split("color:")[1].split(";")[0].strip()
                            run.font.color.rgb = RGBColor(int(color_value[1:3], 16), int(color_value[3:5], 16),
                                                          int(color_value[5:7], 16))

                    elif element.name == "img":
                        src = element.get('src', '')
                        if os.path.isfile(src):
                            paragraph.add_run().add_picture(src)

            if self.is_page_numbering_enabled:
                footer_paragraph = doc.sections[0].footer.paragraphs[0]
                run = footer_paragraph.add_run()

                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                run._r.append(fldChar1)

                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                run._r.append(instrText)

                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar2)

                run.font.size = Pt(12)
                footer_paragraph.alignment = 1

            doc.save(file_path)
            self.statusBar().showMessage(f"Файл сохранен: {file_path}", 5000)
        except Exception as e:
            self.statusBar().showMessage(f"Ошибка сохранения: {e}", 5000)
            print(e)

    def pdf_to_path(self):
        pdf_path, _ = QFileDialog.getSaveFileName(self, "Экспортировать как PDF в...", "", "Документ PDF (*.pdf)")
        docx_path = pdf_path.replace(".pdf", ".docx")
        if docx_path:
            self.save_to_path(docx_path)

        convert(docx_path, pdf_path)
        os.remove(docx_path)

    def convert_to_pdf(self, docx_path, pdf_path):
        try:
            convert(docx_path, pdf_path)
            print(f"PDF сохранён: {pdf_path}")
            return True
        except Exception as e:
            print(f"Ошибка при конвертации в PDF: {e}")
            return False

    def open_document(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Открыть документ...", "", "Документ Word (*.docx)")
        if file_path:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf_file:
                    temp_pdf_path = temp_pdf_file.name

                if not self.convert_to_pdf(file_path, temp_pdf_path):
                    raise Exception("Не удалось конвертировать файл в PDF.")

                reader = PdfReader(temp_pdf_path)
                page_count = len(reader.pages)

                os.remove(temp_pdf_path)

                doc = Document(file_path)
                html_content = ""
                for paragraph in doc.paragraphs:
                    html_content += "<p>"
                    for run in paragraph.runs:
                        style = ""
                        if run.font.size:
                            size = run.font.size.pt
                            style += f"font-size: {size}pt; "
                        if run.font.name:
                            style += f"font-family: {run.font.name}; "
                        if run.font.bold:
                            style += "font-weight: bold; "
                        if run.font.italic:
                            style += "font-style: italic; "
                        if run.font.underline:
                            style += "text-decoration: underline; "
                        if run.font.color.rgb:
                            rgb = run.font.color.rgb
                            style += f"color: #{rgb}; "
                        html_content += f"<span style='{style}'>{run.text}</span>"

                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            image_data = rel.target_part.blob
                            encoded_image = base64.b64encode(image_data).decode('utf-8')
                            html_content += f"<img src='data:image/png;base64,{encoded_image}' style='display:block; margin: 10px 0;' />"

                    html_content += "</p>"

                for sheet in self.sheet.iterate_sheets():
                    sheet.setPlainText("")

                for i in range(page_count):
                    new_sheet = Sheet(self.sheets_layout)
                    print(self.sheets_layout)
                    self.apply_margins_to_stylesheet([0, 0, 0, 0])
                    new_sheet.setHtml(html_content)

                self.current_file_path = file_path
                self.statusBar().showMessage(f"Файл открыт: {file_path}", 5000)

            except Exception as e:
                self.statusBar().showMessage(f"Ошибка открытия: {e}", 5000)
                QMessageBox.critical(self, "Ошибка", f"Не удалось открыть файл: {e}")

    def apply_margins_to_stylesheet(self, margins_in_pixels):
        left_margin, right_margin, top_margin, bottom_margin = margins_in_pixels
        self.sheet.setStyleSheet(f"""
                QTextEdit {{
                    background-color: #FFFFFF;
                    padding-left: {left_margin}px;
                    padding-right: {right_margin}px;
                    padding-top: {top_margin}px;
                    padding-bottom: {bottom_margin}px;
                }}
            """)


app = QApplication(sys.argv)
ex = WordProcessor()
ex.show()
sys.exit(app.exec_())
